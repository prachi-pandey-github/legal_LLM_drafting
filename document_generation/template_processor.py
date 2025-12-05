"""
Template processor for docx-template based document generation
"""
from docxtpl import DocxTemplate
import tempfile
import os
import logging
from typing import Dict, Any
import json

logger = logging.getLogger(__name__)

class TemplateProcessor:
    """Processor for template-based document generation"""
    
    def __init__(self, template_dir: str = "templates"):
        self.template_dir = template_dir
        self.templates = self.load_templates()
    
    def load_templates(self) -> Dict[str, str]:
        """Load available templates"""
        templates = {}
        
        if not os.path.exists(self.template_dir):
            logger.warning(f"Template directory not found: {self.template_dir}")
            return templates
        
        for filename in os.listdir(self.template_dir):
            if filename.endswith('.docx'):
                template_name = filename.replace('.docx', '')
                templates[template_name] = os.path.join(self.template_dir, filename)
                logger.info(f"Loaded template: {template_name}")
        
        return templates
    
    def generate_from_template(self, template_name: str, context: Dict[str, Any]) -> str:
        """Generate document from template"""
        if template_name not in self.templates:
            raise ValueError(f"Template '{template_name}' not found")
        
        try:
            # Load template
            template_path = self.templates[template_name]
            doc = DocxTemplate(template_path)
            
            # Prepare context
            prepared_context = self.prepare_context(context)
            
            # Render template
            doc.render(prepared_context)
            
            # Save to temp file
            temp_dir = tempfile.gettempdir()
            output_filename = f"{template_name}_{tempfile.mktemp(suffix='.docx')}"
            output_path = os.path.join(temp_dir, output_filename)
            
            doc.save(output_path)
            logger.info(f"Generated document from template: {template_name}")
            
            return output_path
            
        except Exception as e:
            logger.error(f"Template generation failed: {str(e)}")
            raise
    
    def prepare_context(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """Prepare context for template rendering"""
        prepared = context.copy()
        
        # Ensure all values are template-friendly
        for key, value in context.items():
            if value is None:
                prepared[key] = ""
            elif isinstance(value, dict) or isinstance(value, list):
                prepared[key] = json.dumps(value, ensure_ascii=False)
            elif isinstance(value, bool):
                prepared[key] = "Yes" if value else "No"
        
        # Add common legal placeholders
        prepared.update({
            "date": "__________________",
            "signature": "__________________",
            "place": "__________________",
            "witness": "__________________",
        })
        
        return prepared
    
    def create_template(self, template_name: str, document_structure: Dict[str, Any]) -> str:
        """Create a new template from document structure"""
        # This is a simplified version
        # In production, you would create an actual .docx template with placeholders
        
        template_path = os.path.join(self.template_dir, f"{template_name}.docx")
        
        # Create basic template
        from docx import Document
        doc = Document()
        
        # Add template structure
        doc.add_heading("{{ title }}", level=0)
        
        for section in document_structure.get("sections", []):
            doc.add_heading(f"{{{{ {section.get('title', '').lower().replace(' ', '_')}_title }}}}", level=1)
            doc.add_paragraph(f"{{{{ {section.get('title', '').lower().replace(' ', '_')}_content }}}}")
        
        # Add signature blocks
        doc.add_heading("Signatures", level=1)
        doc.add_paragraph("Party 1: {{ party1_name }}")
        doc.add_paragraph("Signature: __________________")
        doc.add_paragraph("Date: {{ date }}")
        
        doc.add_paragraph()
        doc.add_paragraph("Party 2: {{ party2_name }}")
        doc.add_paragraph("Signature: __________________")
        doc.add_paragraph("Date: {{ date }}")
        
        # Save template
        os.makedirs(self.template_dir, exist_ok=True)
        doc.save(template_path)
        
        logger.info(f"Created template: {template_path}")
        return template_path