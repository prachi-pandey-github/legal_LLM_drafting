# document_generation/docx_builder.py
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os

class DocxBuilder:
    def __init__(self):
        self.doc = Document()
        # Set compact formatting for single page
        self._set_compact_formatting()
        
    def _set_compact_formatting(self):
        """Set margins and spacing for compact single-page layout"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
    
    def add_title(self, title):
        heading = self.doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Compact title spacing
        heading.space_after = Pt(6)
        
    def add_heading(self, text, level=1):
        h = self.doc.add_heading(text, level=level)
        # Reduce spacing for compact layout
        h.space_before = Pt(3)
        h.space_after = Pt(3)
        
    def add_paragraph(self, text, bold=False, italic=False):
        p = self.doc.add_paragraph(text)
        # Ultra-compact paragraph spacing
        p.space_after = Pt(1)
        if bold:
            p.runs[0].bold = True
        if italic:
            p.runs[0].italic = True
            
    def add_clause(self, clause_title, clause_body):
        self.add_heading(clause_title, level=2)
        self.add_paragraph(clause_body)
        
    def add_signature_block(self, party_name, role):
        # Ultra-compact signature format
        p = self.doc.add_paragraph(f"____________________    {party_name} ({role})")
        p.space_after = Pt(0)
        
    def save_to_temp(self, filename="legal_document.docx"):
        temp_dir = tempfile.gettempdir()
        filepath = os.path.join(temp_dir, filename)
        self.doc.save(filepath)
        return filepath